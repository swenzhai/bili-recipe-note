from __future__ import annotations

import html
import json
import re
import sqlite3
import zipfile
from pathlib import Path


IMAGE_LINE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")


def _plain_lines(markdown: str) -> list[str]:
    lines: list[str] = []
    for raw in markdown.splitlines():
        line = re.sub(r"!\[[^\]]*\]\([^)]+\)", "", raw)
        line = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)
        line = line.lstrip("#").strip()
        line = line.removeprefix("- ").strip()
        if line:
            lines.append(line)
    return lines


def export_obsidian(note_path: Path, output_path: Path | None = None) -> Path:
    output = output_path or note_path.with_name("note.obsidian.md")
    markdown = note_path.read_text(encoding="utf-8")
    title = _plain_lines(markdown)[0] if _plain_lines(markdown) else note_path.parent.name
    content = f"---\ntitle: {json.dumps(title, ensure_ascii=False)}\nsource: bili-recipe-notes\n---\n\n{markdown}"
    output.write_text(content, encoding="utf-8")
    return output


def export_pdf(note_path: Path, output_path: Path | None = None) -> Path:
    output = output_path or note_path.with_suffix(".pdf")
    markdown = note_path.read_text(encoding="utf-8")
    lines = _plain_lines(markdown)
    try:
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import Image, ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        styles = getSampleStyleSheet()
        body = ParagraphStyle(
            "RecipeBody",
            parent=styles["BodyText"],
            fontName="STSong-Light",
            fontSize=10.5,
            leading=16,
            alignment=TA_LEFT,
            spaceAfter=4,
        )
        heading_styles = {
            level: ParagraphStyle(
                f"RecipeHeading{level}",
                parent=body,
                fontSize=max(11, 20 - level * 2),
                leading=max(16, 24 - level * 2),
                spaceBefore=10,
                spaceAfter=6,
            )
            for level in range(1, 7)
        }
        document = SimpleDocTemplate(
            str(output),
            pagesize=A4,
            rightMargin=18 * mm,
            leftMargin=18 * mm,
            topMargin=16 * mm,
            bottomMargin=16 * mm,
            title=note_path.stem,
        )
        story = []
        for raw in markdown.splitlines():
            line = raw.strip()
            if not line:
                story.append(Spacer(1, 3 * mm))
                continue
            image_match = IMAGE_LINE_RE.match(line)
            if image_match:
                image_path = _resolve_local_image(note_path, image_match.group(2))
                if image_path:
                    picture = Image(str(image_path))
                    max_width, max_height = 170 * mm, 105 * mm
                    scale = min(max_width / picture.imageWidth, max_height / picture.imageHeight, 1.0)
                    picture.drawWidth = picture.imageWidth * scale
                    picture.drawHeight = picture.imageHeight * scale
                    story.extend([picture, Spacer(1, 3 * mm)])
                continue
            heading_match = HEADING_RE.match(line)
            if heading_match:
                level = len(heading_match.group(1))
                story.append(Paragraph(html.escape(_inline_plain(heading_match.group(2))), heading_styles[level]))
                continue
            if line.startswith("- "):
                story.append(
                    ListFlowable(
                        [ListItem(Paragraph(html.escape(_inline_plain(line[2:])), body))],
                        bulletType="bullet",
                        leftIndent=14,
                    )
                )
                continue
            story.append(Paragraph(html.escape(_inline_plain(line)), body))
        document.build(story)
        return output
    except Exception:
        _write_basic_pdf(output, lines)
        return output


def export_docx(note_path: Path, output_path: Path | None = None) -> Path:
    output = output_path or note_path.with_suffix(".docx")
    markdown = note_path.read_text(encoding="utf-8")
    lines = _plain_lines(markdown)
    try:
        from docx import Document

        from docx.shared import Inches

        doc = Document()
        for raw in markdown.splitlines():
            line = raw.strip()
            if not line:
                continue
            image_match = IMAGE_LINE_RE.match(line)
            if image_match:
                image_path = _resolve_local_image(note_path, image_match.group(2))
                if image_path:
                    doc.add_picture(str(image_path), width=Inches(5.8))
                    if image_match.group(1):
                        doc.add_paragraph(image_match.group(1), style="Caption")
                continue
            heading_match = HEADING_RE.match(line)
            if heading_match:
                doc.add_heading(_inline_plain(heading_match.group(2)), level=min(len(heading_match.group(1)), 9))
            elif line.startswith("- "):
                doc.add_paragraph(_inline_plain(line[2:]), style="List Bullet")
            else:
                doc.add_paragraph(_inline_plain(line))
        doc.save(output)
        return output
    except Exception:
        _write_basic_docx(output, lines)
        return output


def export_note(note_path: Path, kind: str) -> Path:
    if kind == "pdf":
        return export_pdf(note_path)
    if kind == "docx":
        return export_docx(note_path)
    if kind == "obsidian":
        return export_obsidian(note_path)
    if kind == "zip":
        return export_recipe_bundle(note_path)
    raise ValueError(f"Unsupported export kind: {kind}")


def _inline_plain(markdown: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", markdown)
    text = re.sub(r"[*_`~]+", "", text)
    return text


def _resolve_local_image(note_path: Path, raw_path: str) -> Path | None:
    cleaned = raw_path.strip()
    if not cleaned or "://" in cleaned:
        return None
    folder = note_path.parent.resolve()
    candidate = (folder / cleaned).resolve()
    try:
        candidate.relative_to(folder)
    except ValueError:
        return None
    return candidate if candidate.is_file() else None


def export_recipe_bundle(note_path: Path, output_path: Path | None = None) -> Path:
    folder = note_path.parent
    output = output_path or folder / f"{folder.name}.recipe.zip"
    allowed_names = {
        "note.md",
        "recipe.json",
        "transcript.json",
        "quality.json",
        "job.json",
        "extra_analysis.md",
        "extra_analysis.json",
        "sync-meta.json",
    }
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(folder.rglob("*")):
            if not path.is_file() or path == output:
                continue
            relative = path.relative_to(folder)
            if relative.parts[0] == "images" or path.name in allowed_names:
                archive.write(path, relative.as_posix())
        _add_practice_logs_to_bundle(archive, folder)
    return output


def _add_practice_logs_to_bundle(archive: zipfile.ZipFile, folder: Path) -> None:
    metadata_path = folder / "sync-meta.json"
    if not metadata_path.is_file():
        return
    try:
        recipe_id = str(json.loads(metadata_path.read_text(encoding="utf-8")).get("recipe_id") or "")
    except (OSError, UnicodeError, json.JSONDecodeError):
        return
    database = next(
        (
            parent / ".bili-recipe-notes" / "mobile-sync.sqlite3"
            for parent in [folder.parent, *folder.parents]
            if (parent / ".bili-recipe-notes" / "mobile-sync.sqlite3").is_file()
        ),
        None,
    )
    if not recipe_id or database is None:
        return
    try:
        connection = sqlite3.connect(database)
        connection.row_factory = sqlite3.Row
        rows = connection.execute(
            "SELECT * FROM practice_logs WHERE recipe_id=? ORDER BY cooked_on,created_at", (recipe_id,)
        ).fetchall()
        logs = [dict(row) for row in rows]
        archive.writestr("practice-logs.json", json.dumps(logs, ensure_ascii=False, indent=2) + "\n")
        for digest in {str(row["photo_sha256"]) for row in rows if row["photo_sha256"]}:
            asset = connection.execute("SELECT path FROM assets WHERE sha256=?", (digest,)).fetchone()
            if asset and Path(str(asset["path"])).is_file():
                source = Path(str(asset["path"]))
                archive.write(source, f"practice-images/{source.name}")
        connection.close()
    except (OSError, sqlite3.Error):
        return


def _pdf_hex(text: str) -> str:
    return "<FEFF" + text.encode("utf-16-be", errors="ignore").hex().upper() + ">"


def _write_basic_pdf(output: Path, lines: list[str]) -> None:
    wrapped_lines: list[str] = []
    for line in lines:
        wrapped_lines.extend(line[index : index + 42] for index in range(0, max(1, len(line)), 42))
    pages = [wrapped_lines[index : index + 38] for index in range(0, max(1, len(wrapped_lines)), 38)] or [[]]
    font_number = 3 + len(pages) * 2
    descendant_number = font_number + 1
    page_numbers = [3 + index * 2 for index in range(len(pages))]
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        f"<< /Type /Pages /Kids [{' '.join(f'{number} 0 R' for number in page_numbers)}] /Count {len(pages)} >>".encode("ascii"),
    ]
    for index, page_lines in enumerate(pages):
        content_number = page_numbers[index] + 1
        text_ops = ["BT", "/F1 11 Tf", "48 790 Td"]
        for line_index, line in enumerate(page_lines):
            if line_index:
                text_ops.append("0 -19 Td")
            text_ops.append(f"{_pdf_hex(line)} Tj")
        text_ops.append("ET")
        stream = "\n".join(text_ops).encode("ascii")
        objects.append(
            f"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 {font_number} 0 R >> >> "
            f"/MediaBox [0 0 595 842] /Contents {content_number} 0 R >>".encode("ascii")
        )
        objects.append(b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream")
    objects.extend(
        [
            f"<< /Type /Font /Subtype /Type0 /BaseFont /STSong-Light /Encoding /UniGB-UCS2-H /DescendantFonts [{descendant_number} 0 R] >>".encode("ascii"),
            b"<< /Type /Font /Subtype /CIDFontType0 /BaseFont /STSong-Light /CIDSystemInfo << /Registry (Adobe) /Ordering (GB1) /Supplement 4 >> >>",
        ]
    )
    body = b"%PDF-1.4\n"
    offsets = [0]
    for idx, obj in enumerate(objects, start=1):
        offsets.append(len(body))
        body += f"{idx} 0 obj\n".encode("ascii") + obj + b"\nendobj\n"
    xref_offset = len(body)
    body += f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("ascii")
    for offset in offsets[1:]:
        body += f"{offset:010d} 00000 n \n".encode("ascii")
    body += (
        f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode(
            "ascii"
        )
    )
    output.write_bytes(body)


def _write_basic_docx(output: Path, lines: list[str]) -> None:
    paragraphs = "\n".join(
        f"<w:p><w:r><w:t>{html.escape(line)}</w:t></w:r></w:p>"
        for line in lines
    )
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:body>{paragraphs}<w:sectPr/></w:body></w:document>"
    )
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>",
        )
        zf.writestr(
            "_rels/.rels",
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
            'Target="word/document.xml"/></Relationships>',
        )
        zf.writestr("word/document.xml", document_xml)
